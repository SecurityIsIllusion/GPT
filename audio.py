import os
import openai
import docx
from docx import Document
document = Document()
openai.organization="org-i*********************sws"
openai.api_key=os.getenv("OPENAI_API_KEY")
def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
   	 transcription = openai.Audio.translate("whisper-1", audio_file)
   	 document.add_heading('Audio Translate', level=1)
   	 document.add_paragraph(transcription["text"])
   	 document.save('demo.docx')
    return transcription['text']

while True:
    user_input=input("Insert Path:  ")
    if user_input == 'exit':
   	 print ('Bye!!!')
   	 break
    response=transcribe_audio(user_input)
    print("ChatGPT: ",response)
