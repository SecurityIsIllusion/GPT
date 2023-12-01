import os
import openai
import docx
from docx import Document
document = Document()
from openai import OpenAI
client = OpenAI()
openai.organization="org-i*********************sws"
openai.api_key=os.getenv("OPENAI_API_KEY")
def transcribe_audio(audio_file_path):
#	with open(audio_file_path, 'rb') as audio_file:
#	audio_file = open("speech.mp3", "rb")
	audio_file = open(audio_file_path, "rb")
#       	transcription = openai.Audio.transcribe("whisper-1", audio_file)
	transcript = client.audio.transcriptions.create(
		model="whisper-1", 
		file=audio_file, 
  		response_format="text"
  		)
	document.add_heading('Audio Translate', level=1)
#	document.add_paragraph(transcript["text"])
	document.add_paragraph(transcript)
	document.save('demo.docx')
#	return transcript['text']
	return transcript

while True:
	user_input=input("Insert Path:  ")
	if user_input == 'exit':
		print ('Bye!!!')
		break
	response=transcribe_audio(user_input)
	print("ChatGPT: ",response)
